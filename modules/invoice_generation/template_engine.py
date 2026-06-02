from docx import Document


def replace_placeholders(doc_path, output_path, data):
    doc = Document(doc_path)

    # Replace placeholders in normal paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = "{{" + key + "}}"

            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    placeholder,
                    str(value)
                )

    # Replace placeholders inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = "{{" + key + "}}"

                    if placeholder in cell.text:
                        cell.text = cell.text.replace(
                            placeholder,
                            str(value)
                        )

    doc.save(output_path)